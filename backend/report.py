"""STR (Suspicious Transaction Report) generation — PDF via reportlab.

Bundles everything the fused analysis found into a police-readable PDF:
case summary, entity table, risk-ranked accounts and phones, money-flow
highlights, temporal coincidence windows, and payout patterns.

DOCX (editable Word) output is provided by generate_docx_report() for
investigators who need to annotate the report before filing it.
"""

from __future__ import annotations

import os
import re
from datetime import datetime

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (Paragraph, SimpleDocTemplate, Spacer, Table,
                                TableStyle)

from .fusion import correlate_phones, fraud_heat, rapid_payouts
from .graphs import summary_graphs

_DARK = colors.HexColor("#1f2d3d")
_ACCENT = colors.HexColor("#c0392b")
_GREY = colors.HexColor("#5b6b7b")


def _h(styles, text):
    p = Paragraph(text, styles["Heading2"])
    return p


def _para(text):
    return Paragraph(text)


def _table(headers: list[str], rows: list[list], widths=None) -> Table:
    t = Table([headers] + rows, colWidths=widths, repeatRows=1)
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), _DARK),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 7.5),
        ("GRID", (0, 0), (-1, -1), 0.4, _GREY),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f4f6f8")]),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
    ]))
    return t


def _money(v) -> str:
    return "%.2f" % (v or 0)


def _risk_chart(accounts: list[dict]) -> "Drawing | None":
    """Horizontal bar chart of top composite risk scores (no new deps)."""
    from reportlab.graphics.charts.barcharts import VerticalBarChart
    from reportlab.graphics.shapes import Drawing, String

    top = [a for a in accounts if a["score"] > 0][:10]
    if not top:
        return None
    chart = VerticalBarChart()
    chart.data = [[a["score"] for a in reversed(top)]]
    chart.categoryAxis.categoryNames = [str(a["account_no"])[-8:] for a in reversed(top)]
    chart.valueAxis.valueMin = 0
    chart.valueAxis.valueMax = 100
    chart.valueAxis.labelTextFormat = "%d"
    chart.bars[0].fillColor = _ACCENT
    chart.categoryAxis.labels.fontSize = 5.5
    chart.valueAxis.labels.fontSize = 6
    chart.x = 30
    chart.y = 20
    chart.width = 480
    chart.height = 110
    d = Drawing(540, 160)
    d.add(chart)
    d.add(String(30, 140, "Composite risk score by account (top 10)",
                 fontSize=8, textColor=_DARK))
    return d


def generate_str_report(bundle: dict, out_path: str, case_title: str = "") -> str:
    heat = fraud_heat(bundle)
    hits = correlate_phones(bundle)
    rapids = rapid_payouts(bundle)
    graphs = summary_graphs(bundle)
    complaints = bundle.get("complaints", [])
    bank = bundle.get("bank", [])
    cdr = bundle.get("cdr", [])
    ipdr = bundle.get("ipdr", [])

    doc = SimpleDocTemplate(out_path, pagesize=A4,
                            rightMargin=14 * mm, leftMargin=14 * mm,
                            topMargin=14 * mm, bottomMargin=14 * mm)
    ss = getSampleStyleSheet()
    title_style = ParagraphStyle("Title", parent=ss["Title"], fontSize=16,
                                 textColor=_DARK, spaceAfter=2)
    sub_style = ParagraphStyle("Sub", parent=ss["Normal"], fontSize=9,
                               textColor=_GREY, spaceAfter=10)
    h2 = ParagraphStyle("H2", parent=ss["Heading2"], fontSize=11,
                        textColor=_ACCENT, spaceBefore=12, spaceAfter=4)
    tiny = ParagraphStyle("Tiny", parent=ss["Normal"], fontSize=7)

    el = []
    el.append(Paragraph("Suspicious Transaction Report", title_style))
    el.append(Paragraph(
        f"Generated {datetime.now().strftime('%Y-%m-%d %H:%M')} &nbsp;|&nbsp; "
        f"{case_title or 'AI-assisted analysis'} "
        f"&nbsp;|&nbsp; {len(bank)} bank txns, {len(cdr)} CDR records, "
        f"{len(ipdr)} IPDR sessions, {len(complaints)} NCRP complaints", sub_style))

    total_in = sum(r.get("credit") or 0 for r in bank)
    total_out = sum(r.get("debit") or 0 for r in bank)
    el.append(_h(ss, "1. Executive summary"))
    el.append(_para(
        f"Analysed <b>{len(bank)}</b> bank transactions across "
        f"<b>{len(graphs['top_accounts'])}</b> accounts, "
        f"<b>{len(cdr)}</b> CDR records involving <b>{graphs['phone_call_graph']['nodes']}</b> "
        f"unique phone numbers and <b>{len(ipdr)}</b> internet sessions. "
        f"Total credits observed: <b>Rs {_money(total_in)}</b>; total debits: "
        f"<b>Rs {_money(total_out)}</b>. "
        f"<b>{len([a for a in heat['accounts'] if a['score'] >= 50])}</b> accounts carry "
        f"high composite risk scores and are detailed below."))

    el.append(_h(ss, "2. Accounts by risk"))
    rows = []
    for a in heat["accounts"]:
        rows.append([
            a["account_no"], a["bank"], a["txns"], _money(a["credit"]),
            _money(a["debit"]), f"{a['score']}/100", ", ".join(a["flags"])[:80],
        ])
    if rows:
        el.append(_table(["Account", "Bank", "Txns", "Credits Rs", "Debits Rs",
                          "Risk", "Flags"], rows))
    else:
        el.append(_para("No accounts found."))
    chart = _risk_chart(heat["accounts"])
    if chart is not None:
        el.append(Spacer(1, 8))
        el.append(chart)

    el.append(_h(ss, "3. Phones by activity / risk"))
    rows = []
    for p in heat["phones"]:
        rows.append([p["phone"], p["records"], p["contacts"],
                     p["unique_contacts"], p["sms"], p["voice"],
                     f"{p['score']}/100", ", ".join(p["flags"])[:60]])
    if rows:
        el.append(_table(["Phone", "CDR recs", "Contacts", "Unique", "SMS",
                          "Voice", "Risk", "Flags"], rows[:25]))
    else:
        el.append(_para("No CDR records."))

    el.append(_h(ss, "4. Phone call network"))
    el.append(_para(
        f"Call graph: <b>{graphs['phone_call_graph']['nodes']}</b> nodes / "
        f"<b>{graphs['phone_call_graph']['edges']}</b> edges. Most connected targets:"))
    rows = [[c["phone"], c["degree"], c["out"], c["calls"]]
            for c in graphs["central_phones"][:10]]
    if rows:
        el.append(_table(["Phone", "Degree", "Outgoing", "Calls"], rows))

    el.append(_h(ss, "5. Bank <-> telecom coincidence"))
    if hits["hits"]:
        rows = []
        for h in hits["hits"][:30]:
            rows.append([
                h["phone"], h["account_no"], h["txn_date"], h["mode"],
                _money(h["amount"]), h["phone_cdr_records"], h["window_count"],
            ])
        el.append(_table(["Phone", "Account", "Txn date", "Mode", "Amount Rs",
                          "CDR recs", "In-window"], rows))
    else:
        el.append(_para("No direct phone overlap between bank counterparties "
                        "and CDR subscribers in this bundle."))

    el.append(_h(ss, "6. Payout patterns"))
    rows = [[r["account_no"], r["count"], r["window_min"], _money(r["total"]),
             f"{datetime.fromtimestamp(r['start_ts']).strftime('%Y-%m-%d %H:%M')} "
             f"to {datetime.fromtimestamp(r['end_ts']).strftime('%Y-%m-%d %H:%M')}"]
            for r in rapids[:12]]
    if rows:
        el.append(_table(["Account", "Payouts", "Window(min)", "Total Rs",
                          "Period"], rows))
    rows = [[x["txn_id"][:34], x["account_no"], x["date"], x["mode"],
             _money(x["amount"]), (x["narration"] or "")[:60]]
            for x in heat["round_payouts"][:20]]
    if rows:
        el.append(Spacer(1, 6))
        el.append(_para("Round-amount debit transactions (Rs 5k+ multiples of 5000):"))
        el.append(Spacer(1, 4))
        el.append(_table(["Txn id", "Account", "Date", "Mode", "Amount Rs",
                          "Narration"], rows))

    el.append(_h(ss, "7. NCRP fraud-account complaints"))
    if complaints:
        states = sorted({c.get("state") for c in complaints if c.get("state")})
        accts = sorted({c.get("account_no") for c in complaints})
        el.append(_para(
            f"<b>{len(complaints)}</b> NCRP complaint rows referencing "
            f"<b>{len(accts)}</b> beneficiary accounts across "
            f"<b>{len(states)}</b> states. Accounts named in complaints are "
            f"auto-flagged in section 2."))
    else:
        el.append(_para("No NCRP complaint ledger ingested."))

    doc.build(el)
    return out_path


def generate_entity_str_report(bundle: dict, kind: str, value: str,
                               out_path: str) -> str:
    """Individual STR PDF for one entity (node in the investigation graph)."""
    from .evidence import entity_intelligence

    info = entity_intelligence(bundle, kind, value)
    if info is None:
        raise ValueError(f"no evidence for {kind} {value}")

    doc = SimpleDocTemplate(out_path, pagesize=A4,
                            rightMargin=14 * mm, leftMargin=14 * mm,
                            topMargin=14 * mm, bottomMargin=14 * mm)
    ss = getSampleStyleSheet()
    title_style = ParagraphStyle("Title", parent=ss["Title"], fontSize=16,
                                 textColor=_DARK, spaceAfter=2)
    sub_style = ParagraphStyle("Sub", parent=ss["Normal"], fontSize=9,
                               textColor=_GREY, spaceAfter=10)
    h2 = ParagraphStyle("H2", parent=ss["Heading2"], fontSize=11,
                        textColor=_ACCENT, spaceBefore=12, spaceAfter=4)
    mono = ParagraphStyle("Mono", parent=ss["Normal"], fontName="Courier",
                          fontSize=8.5, spaceAfter=6)

    el = []
    el.append(Paragraph("Suspicious Transaction Report", title_style))
    el.append(Paragraph(
        f"Entity: <b>{kind.upper()}</b> · <font name='Courier'>{value}</font>",
        sub_style))
    el.append(Paragraph(
        f"Generated {datetime.now().strftime('%Y-%m-%d %H:%M')} | "
        f"composite risk {info['risk_score']}/100 ({info['risk_band']}) | "
        f"confidence {info['confidence']:.0%}", sub_style))

    el.append(_h(ss, "1. Executive summary"))
    v = info["volumes"]
    c = info["counts"]
    el.append(_para(
        f"<b>{value}</b> is a {kind.upper()} entity with "
        f"<b>{c['transactions']}</b> bank transaction(s), "
        f"<b>{c['calls']}</b> call(s), <b>{c['sms']}</b> SMS and "
        f"<b>{c['ip_sessions']}</b> IP session(s). "
        f"Credits Rs <b>{_money(v['credit'])}</b>, debits Rs "
        f"<b>{_money(v['debit'])}</b>, average Rs <b>{_money(v['avg_amount'])}</b>, "
        f"largest Rs <b>{_money(v['max_amount'])}</b>. "
        f"Activity: <b>{info['activity']['first'] or '—'}</b> → "
        f"<b>{info['activity']['last'] or '—'}</b>. "
        f"Composite risk score <b>{info['risk_score']}/100</b> — "
        f"<b>{info['risk_band']}</b> band, confidence {info['confidence']:.0%}."))

    el.append(_h(ss, "2. Risk explanation (why this score?)"))
    bd = info.get("breakdown") or []
    if bd:
        rows = [[x.get("rule", ""), f"+{x.get('points', 0)}",
                 x.get("reason", "")] for x in bd]
        el.append(_table(["Rule", "Points", "Reason"], rows))
    else:
        el.append(_para("No rule-based contributions — score is low/zero."))
    if info.get("flags"):
        el.append(_para("Flags: <b>" + ", ".join(info["flags"]) + "</b>"))

    el.append(_h(ss, "3. Suspicious patterns detected"))
    pats = info.get("patterns") or []
    if pats:
        rows = [[p.get("label", ""), p.get("evidence", "")] for p in pats]
        el.append(_table(["Pattern", "Evidence"], rows))
    else:
        el.append(_para("No anomalous patterns detected for this entity."))

    el.append(_h(ss, "4. Linked entities"))
    links = info.get("links") or {}
    if links:
        rows = []
        for name, items in links.items():
            if items:
                rows.append([name, ", ".join(str(x) for x in items[:20])])
        if rows:
            el.append(_table(["Relationship", "Entities"], rows))
    else:
        el.append(_para("No linked entities found."))
    if info.get("ncrp"):
        el.append(_para(f"NCRP complaints: <b>{len(info['ncrp'])}</b> "
                        f"ledger row(s) reference this entity."))

    el.append(_h(ss, "5. Recent evidence records"))
    recs = info.get("records") or []
    if recs:
        rows = [[r.get("kind", ""), f"{r.get('date') or ''} {r.get('time') or ''}",
                 (r.get("label") or "")[:90],
                 _money(r.get("amount")) if r.get("amount") else ""]
                for r in recs]
        el.append(_table(["Type", "When", "Detail", "Rs"], rows))
    else:
        el.append(_para("No recent records."))

    el.append(_h(ss, "6. Recommended investigation actions"))
    rec = []
    if info["risk_score"] >= 50:
        rec.append("File individual STR; freeze the entity pending review.")
    if info["risk_score"] >= 25:
        rec.append("Request banking / telecom records for the linked entities "
                   "in section 4.")
    if v.get("round_amounts"):
        rec.append("Examine round-amount payouts for structuring indicators.")
    if any(p.get("label") in ("RAPID IN-AND-OUT (mule signature)",
                              "RAPID CASH-OUT") for p in pats):
        rec.append("Prioritise mule-account cash-through analysis; "
                   "map onward beneficiaries.")
    if any(p.get("label") == "CIRCULAR FLOW" for p in pats):
        rec.append("Trace circular-flow cycle participants and their "
                   "account holders.")
    if any(p.get("label") == "SHARED DEVICE" for p in pats):
        rec.append("Pull IMEI tower data for co-users of the shared device.")
    if not rec:
        rec.append("No immediate action required; continue routine monitoring.")
    for i, r in enumerate(rec, 1):
        el.append(_para(f"{i}. {r}"))

    doc.build(el)
    return out_path


def generate_transaction_str_report(bundle: dict, txn_id: str,
                                    out_path: str) -> str:
    """Focused STR for a single transaction: identity, hybrid risk
    decomposition, named scenarios, rule evidence, timeline, money-flow leg
    and investigator recommendations (Hybrid Fraud Detection Engine)."""
    from .risk.engine import transaction_risk
    from .risk.hybrid import (explanations_for_txn, hybrid_analyze,
                              hybrid_transaction_risk)

    bank = bundle.get("bank", [])
    txn = next((r for r in bank
                if (r.get("txn_id") or r.get("transaction_id")) == txn_id), None)
    if txn is None:
        raise ValueError(f"transaction {txn_id} not in bundle")

    scored = {s["transaction_id"]: s
              for s in transaction_risk(bundle)}
    s = scored.get(txn_id, {})
    comp = s.get("risk_components", {})
    phone = txn.get("sender_phone") or txn.get("receiver_phone") or ""

    # Hybrid engine artefacts (scenarios, model scores, explanation).
    try:
        hybrid_rows = {r["transaction_id"]: r
                       for r in hybrid_transaction_risk(bundle)}
        hrec = hybrid_rows.get(txn_id, {})
        hscen = hrec.get("scenarios") or []
        hcomps = hrec.get("hybrid_components") or {}
        hmodels = hrec.get("models_fired") or []
        hexpl = explanations_for_txn(bundle, txn_id)
        htiles = [e for e in hexpl.get("timeline", [])][:8]
        hrecs = hexpl.get("recommendations") or []
    except Exception:  # noqa: BLE001 — hybrid engine is best-effort in STR
        hrec, hscen, hcomps, hmodels, hexpl, htiles, hrecs = \
            {}, [], {}, [], {}, [], []

    doc = SimpleDocTemplate(out_path, pagesize=A4,
                            rightMargin=14 * mm, leftMargin=14 * mm,
                            topMargin=14 * mm, bottomMargin=14 * mm)
    ss = getSampleStyleSheet()
    title_style = ParagraphStyle("Title", parent=ss["Title"], fontSize=16,
                                 textColor=_DARK, spaceAfter=2)
    sub_style = ParagraphStyle("Sub", parent=ss["Normal"], fontSize=9,
                               textColor=_GREY, spaceAfter=10)
    h2 = ParagraphStyle("H2", parent=ss["Heading2"], fontSize=11,
                        textColor=_ACCENT, spaceBefore=12, spaceAfter=4)

    el = []
    el.append(Paragraph("Transaction STR — Suspicious Transaction Report",
                        title_style))
    el.append(Paragraph(
        f"Transaction <b>{txn_id}</b> &nbsp;|&nbsp; generated "
        f"{datetime.now().strftime('%Y-%m-%d %H:%M')}", sub_style))

    el.append(_h(ss, "1. Transaction identity"))
    el.append(_table(
        ["Field", "Value"],
        [["Transaction ID", txn_id],
         ["Account", txn.get("account_no") or ""],
         ["Customer", txn.get("customer_id") or txn.get("sender_customer_id") or ""],
         ["Amount", f"Rs {_money(txn.get('credit') or txn.get('debit') or 0)}"],
         ["Mode", txn.get("mode") or ""],
         ["Date / time", f"{txn.get('date') or ''} {txn.get('time') or ''}".strip()],
         ["Counterparty", txn.get("counterparty_name") or ""],
         ["Receiver account", txn.get("receiver_account") or ""],
         ["Phone", phone],
         ["Narration", str(txn.get("narration") or "")[:120]]],
        widths=[45 * mm, 120 * mm]))

    el.append(_h(ss, "2. Hybrid risk decomposition"))
    el.append(_table(
        ["Source", "Score"],
        [["Behavioural rules", f"{comp.get('behavioural', 0):.1f}"],
         ["Txn-level ML", f"{comp.get('txn_ml', 0):.1f}"],
         ["Account composite", f"{comp.get('account_composite', 0):.1f}"],
         ["Profile deviation", f"{hcomps.get('behaviour', 0):.1f}"],
         ["Temporal correlation", f"{hcomps.get('temporal', 0):.1f}"],
         ["Telecom context", f"{hcomps.get('telecom', 0):.1f}"],
         ["Internet context", f"{hcomps.get('internet', 0):.1f}"],
         ["Composite risk",
          f"{s.get('risk_score', 0):.1f} ({s.get('risk_band', 'SAFE')})"]],
        widths=[45 * mm, 120 * mm]))

    if hmodels:
        el.append(_h(ss, "3. Engines triggered"))
        el.append(_para("Detectors that fired: "
                        + ", ".join(f"<b>{m}</b>" for m in hmodels) + "."))

    if hscen:
        el.append(_h(ss, "4. Fraud scenarios"))
        for sc in hscen:
            el.append(_para(
                f"<b>{sc['scenario']}</b> — {sc.get('description', '')} "
                f"(severity {sc.get('severity', '')}, "
                f"confidence {float(sc.get('confidence', 0)):.0%})"))
            for ev in sc.get("evidence", [])[:4]:
                el.append(_para(f"&nbsp;&nbsp;• {ev}"))

    rules = s.get("breakdown") or []
    if rules:
        el.append(_h(ss, "5. Rules fired"))
        el.append(_table(
            ["Rule", "Pts", "W", "Reason"],
            [[r.get("rule"), r.get("points"), r.get("weight"),
              str(r.get("reason") or "")[:80]] for r in rules],
            widths=[48 * mm, 14 * mm, 12 * mm, 91 * mm]))

    ev = s.get("evidence") or []
    if ev:
        el.append(_h(ss, "6. Evidence"))
        for i, e in enumerate(ev, 1):
            el.append(_para(f"{i}. {e}"))

    if hexpl.get("narrative"):
        el.append(_h(ss, "7. Investigative summary"))
        el.append(_para(hexpl["narrative"]))

    if htiles:
        el.append(_h(ss, "8. Activity timeline (±1 hour)"))
        for e in htiles[:6]:
            el.append(_para(f"{e.get('kind', '')}: {e.get('detail', '')}"))

    receiver = txn.get("receiver_account") or ""
    if receiver:
        el.append(_h(ss, "9. Receiver leg"))
        inflow = sum(float(r.get("credit") or 0.0) for r in bank
                     if (r.get("receiver_account") or "") == receiver)
        outflows = [r for r in bank if (r.get("receiver_account") or "") == receiver]
        el.append(_para(
            f"Receiver <b>{receiver}</b> appears in "
            f"<b>{len(outflows)}</b> transactions of this bundle "
            f"(total inflows Rs {_money(inflow)})."))

    if hrecs:
        el.append(_h(ss, "10. Recommendations"))
        for i, r in enumerate(hrecs, 1):
            el.append(_para(f"{i}. {r}"))

    doc.build(el)
    return out_path


def generate_docx_report(bundle: dict, out_path: str,
                         case_title: str = "") -> str:
    """Forensic report as an editable Word document (python-docx)."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    heat = fraud_heat(bundle)
    hits = correlate_phones(bundle)
    rapids = rapid_payouts(bundle)
    complaints = bundle.get("complaints", [])
    bank = bundle.get("bank", [])
    cdr = bundle.get("cdr", [])
    ipdr = bundle.get("ipdr", [])

    doc = Document()
    dark = RGBColor(0x1F, 0x2D, 0x3D)
    accent = RGBColor(0xC0, 0x39, 0x2B)
    grey = RGBColor(0x5B, 0x6B, 0x7B)

    title = doc.add_heading("Suspicious Transaction Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    run = sub.add_run(
        f"Generated {datetime.now().strftime('%Y-%m-%d %H:%M')} | "
        f"{case_title or 'AI-assisted analysis'} | "
        f"{len(bank)} bank txns, {len(cdr)} CDR records, "
        f"{len(ipdr)} IPDR sessions, {len(complaints)} NCRP complaints")
    run.font.color.rgb = grey
    run.font.size = Pt(9)

    def heading(text):
        h = doc.add_heading(text, level=1)
        for r in h.runs:
            r.font.color.rgb = accent
        return h

    def table(headers, rows):
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Light Grid Accent 1"
        for i, h in enumerate(headers):
            cell = t.rows[0].cells[i]
            cell.text = h
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.bold = True
                    r.font.size = Pt(8)
        for row in rows:
            cells = t.add_row().cells
            for i, v in enumerate(row):
                cells[i].text = str(v)
                for p in cells[i].paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(8)
        return t

    total_in = sum(r.get("credit") or 0 for r in bank)
    total_out = sum(r.get("debit") or 0 for r in bank)
    heading("1. Executive summary")
    doc.add_paragraph(
        f"Analysed {len(bank)} bank transactions across "
        f"{len({r.get('account_no') for r in bank})} accounts, {len(cdr)} CDR "
        f"records and {len(ipdr)} internet sessions. Total credits observed: "
        f"Rs {total_in:,.2f}; total debits: Rs {total_out:,.2f}. "
        f"{len([a for a in heat['accounts'] if a['score'] >= 50])} accounts "
        f"carry high composite risk scores and are detailed below.")

    heading("2. Accounts by risk")
    rows = [[a["account_no"], a["bank"], a["txns"], _money(a["credit"]),
             _money(a["debit"]), f"{a['score']}/100", ", ".join(a["flags"])[:80]]
            for a in heat["accounts"]]
    if rows:
        table(["Account", "Bank", "Txns", "Credits Rs", "Debits Rs", "Risk",
               "Flags"], rows)
    else:
        doc.add_paragraph("No accounts found.")

    heading("3. Phones by activity / risk")
    rows = [[p["phone"], p["records"], p["contacts"], p["unique_contacts"],
             p["sms"], p["voice"], f"{p['score']}/100",
             ", ".join(p["flags"])[:60]] for p in heat["phones"][:25]]
    if rows:
        table(["Phone", "CDR recs", "Contacts", "Unique", "SMS", "Voice",
               "Risk", "Flags"], rows)
    else:
        doc.add_paragraph("No CDR records.")

    heading("4. Bank <-> telecom coincidence")
    if hits["hits"]:
        rows = [[h["phone"], h["account_no"], h["txn_date"], h["mode"],
                 _money(h["amount"]), h["phone_cdr_records"], h["window_count"]]
                for h in hits["hits"][:30]]
        table(["Phone", "Account", "Txn date", "Mode", "Amount Rs",
               "CDR recs", "In-window"], rows)
    else:
        doc.add_paragraph("No direct phone overlap between bank counterparties "
                          "and CDR subscribers in this bundle.")

    heading("5. Payout patterns")
    rows = [[r["account_no"], r["count"], r["window_min"], _money(r["total"]),
             f"{datetime.fromtimestamp(r['start_ts']).strftime('%Y-%m-%d %H:%M')} "
             f"to {datetime.fromtimestamp(r['end_ts']).strftime('%Y-%m-%d %H:%M')}"]
            for r in rapids[:12]]
    if rows:
        table(["Account", "Payouts", "Window(min)", "Total Rs", "Period"], rows)
    if heat["round_payouts"]:
        doc.add_paragraph("Round-amount debit transactions (Rs 5k+ multiples "
                          "of 5000):")
        rows = [[x["txn_id"][:34], x["account_no"], x["date"], x["mode"],
                 _money(x["amount"]), (x["narration"] or "")[:60]]
                for x in heat["round_payouts"][:20]]
        table(["Txn id", "Account", "Date", "Mode", "Amount Rs", "Narration"],
              rows)

    heading("6. NCRP fraud-account complaints")
    if complaints:
        accts = sorted({c.get("account_no") for c in complaints})
        doc.add_paragraph(
            f"{len(complaints)} NCRP complaint rows referencing {len(accts)} "
            f"beneficiary accounts. Accounts named in complaints are "
            f"auto-flagged in section 2.")
    else:
        doc.add_paragraph("No NCRP complaint ledger ingested.")

    doc.save(out_path)
    return out_path
